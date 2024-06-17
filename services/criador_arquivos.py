from services.alunos import *
from docx import Document
from datetime import datetime
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

date_str = datetime.now().strftime("%d-%m-%Y")

alunos_matriculados = {
    "[ALUNOS_MATRICULADOS_CLASSE_ABRAÃO]": count_alunos_turma_abraao(),
    "[ALUNOS_MATRICULADOS_CLASSE_SARAH]": count_alunos_turma_sarah(),
    "[ALUNOS_MATRICULADOS_CLASSE_AMÓS_E_MIRIÃ]": count_alunos_turma_amos(),
    "[ALUNOS_MATRICULADOS_CLASSE_SAMUEL]": count_alunos_turma_samuel(),
}
total_matriculados_geral = sum(alunos_matriculados.values())
alunos_matriculados["[ALUNOS_MATRICULADOS_TOTAL_GERAL]"] = total_matriculados_geral


alunos_ausentes = {
    "[AUSENTES_CLASSE_ABRAÃO]": count_alunos_faltas_turma_abraao(date_str),
    "[AUSENTES_CLASSE_SARAH]": count_alunos_faltas_turma_sarah(date_str),
    "[AUSENTES_CLASSE_AMÓS_E_MIRIÃ]": count_alunos_faltas_turma_amos(date_str),
    "[AUSENTES_CLASSE_SAMUEL]": count_alunos_faltas_turma_samuel(date_str),
}
total_ausentes_geral = sum(alunos_ausentes.values())
alunos_ausentes["[AUSENTES_TOTAL_GERAL]"] = total_ausentes_geral

visitantes = {
    "[VISITANTES_CLASSE_ABRAÃO]": get_visitantes_turma_abraao(date_str),
    "[VISITANTES_CLASSE_SARAH]": get_visitantes_turma_sarah(date_str),
    "[VISITANTES_CLASSE_AMÓS_E_MIRIÃ]": get_visitantes_turma_amos(date_str),
    "[VISITANTES_CLASSE_SAMUEL]": get_visitantes_turma_samuel(date_str),
}
total_visitantes_geral = sum(visitantes.values())
visitantes["[VISITANTES_TOTAL_GERAL]"] = total_visitantes_geral

alunos_presentes = {
    "[PRESENTES_CLASSE_ABRAÃO]": alunos_matriculados["[ALUNOS_MATRICULADOS_CLASSE_ABRAÃO]"] - alunos_ausentes["[AUSENTES_CLASSE_ABRAÃO]"] + visitantes["[VISITANTES_CLASSE_ABRAÃO]"],
    "[PRESENTES_CLASSE_SARAH]": alunos_matriculados["[ALUNOS_MATRICULADOS_CLASSE_SARAH]"] - alunos_ausentes["[AUSENTES_CLASSE_SARAH]"] + visitantes["[VISITANTES_CLASSE_SARAH]"],
    "[PRESENTES_CLASSE_AMÓS_E_MIRIÃ]": alunos_matriculados["[ALUNOS_MATRICULADOS_CLASSE_AMÓS_E_MIRIÃ]"] - alunos_ausentes["[AUSENTES_CLASSE_AMÓS_E_MIRIÃ]"] + visitantes["[VISITANTES_CLASSE_AMÓS_E_MIRIÃ]"],
    "[PRESENTES_CLASSE_SAMUEL]": alunos_matriculados["[ALUNOS_MATRICULADOS_CLASSE_SAMUEL]"] - alunos_ausentes["[AUSENTES_CLASSE_SAMUEL]"] + visitantes["[VISITANTES_CLASSE_SAMUEL]"],
}
total_presentes_geral = sum(alunos_presentes.values())
alunos_presentes["[PRESENTES_TOTAL_GERAL]"] = total_presentes_geral

quantidade_biblias = {
    "[BIBLIAS_CLASSE_ABRAÃO]": get_biblias_by_date_and_turma_abraao(date_str),
    "[BIBLIAS_CLASSE_AMÓS_E_MIRIÃ]": get_biblias_by_date_and_turma_amos(date_str),
    "[BIBLIAS_CLASSE_SARAH]": get_biblias_by_date_and_turma_sarah(date_str),
    "[BIBLIAS_CLASSE_SAMUEL]": get_biblias_by_date_and_turma_samuel(date_str),
}
total_quantidade_biblias = sum(quantidade_biblias.values())
quantidade_biblias["[BIBLIAS_TOTAL_GERAL]"] = total_quantidade_biblias


quantidade_revistas = {
    "[REVISTAS_CLASSE_ABRAÃO]": get_revistas_by_date_and_turma_abraao(date_str),
    "[REVISTAS_CLASSE_AMÓS_E_MIRIÃ]": get_revistas_by_date_and_turma_amos(date_str),
    "[REVISTAS_CLASSE_SARAH]": get_revistas_by_date_and_turma_sarah(date_str),
    "[REVISTAS_CLASSE_SAMUEL]": get_revistas_by_date_and_turma_samuel(date_str),
}
total_quantidade_revistas = sum(quantidade_revistas.values())
quantidade_revistas["[REVISTAS_TOTAL_GERAL]"] = total_quantidade_revistas

quantidade_retardatarios = {
    "[RETARDATARIOS_CLASSE_ABRAÃO]": get_retardatarios_by_date_and_turma_abraao(date_str),
    "[RETARDATARIOS_CLASSE_AMÓS_E_MIRIÃ]": get_retardatarios_by_date_and_turma_amos(date_str),
    "[RETARDATARIOS_CLASSE_SARAH]": get_retardatarios_by_date_and_turma_sarah(date_str),
    "[RETARDATARIOS_CLASSE_SAMUEL]": get_retardatarios_by_date_and_turma_samuel(date_str),
}
total_quantidade_retardatarios = sum(quantidade_retardatarios.values())
quantidade_retardatarios["[RETARDATARIOS_TOTAL_GERAL]"] = total_quantidade_retardatarios

def MontaDocumento(template_path, output_path):
    print("Buscando data: ", date_str)
    document = Document(template_path)
    substitute_date_in_docx_template(document)
    add_title_above_table(document, "RESULTADOS")
    table = document.add_table(rows=8, cols=6)
    CriarTabela(table)
    add_title_above_table(document, "CLASSIFICAÇÃO")
    table2 = document.add_table(rows=5, cols=4)
    ordenar_valores(table2)

    
    # for paragraph in document.paragraphs:
    #     for key, value in tabela_dados.items():
    #         if key in paragraph.text:
    #             for run in paragraph.runs:
    #                 run.text = run.text.replace(key, value)
                    
    document.save(output_path)
    # for key, value in visitantes.items():
    #     print(key, value)
    # for key, value in alunos_presentes.items():
    #     print(key, value)
    
    
    
def add_title_above_table(doc, title_text):
    # Adiciona um título centralizado acima da tabela
    paragraph = doc.add_paragraph()
    paragraph.alignment = 1  # Alinhamento centralizado (0 = esquerda, 1 = centralizado, 2 = direita)
    run = paragraph.add_run(title_text)
    font = run.font
    font.bold = True


def substitute_date_in_docx_template(doc):
    # Obter a data atual
    current_date = datetime.now()

    # Substituir os marcadores de data no documento
    formatted_date = current_date.strftime("%d/%m/%Y")
    for paragraph in doc.paragraphs:
        if "[dia]" in paragraph.text:
            paragraph.text = paragraph.text.replace("[dia]", str(current_date.day))
        if "[mes]" in paragraph.text:
            paragraph.text = paragraph.text.replace("[mes]", str(current_date.month))
        if "[ano]" in paragraph.text:
            paragraph.text = paragraph.text.replace("[ano]", str(current_date.year))

def set_cell_background(cell, color):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def CriarTabela(table):
    for i in range(8):
        for j in range(6):
            cell = table.cell(i, j)
            if j == 0:
                set_cell_background(cell, '1F3863')
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.bold = True
                        
                        
    headers = table.rows[0]
    headers.cells[0].text = ""
    headers.cells[1].text = "Classe Abraão"
    headers.cells[2].text = "Classe Sarah"
    headers.cells[3].text = "Classe Amós e Miriã"
    headers.cells[4].text = "Classe Samuel"
    headers.cells[5].text = "Total Geral"
    
    alunos = table.rows[1]
    alunos.cells[0].text = "Alunos Matriculados"
    alunos.cells[1].text = alunos_matriculados.get("[ALUNOS_MATRICULADOS_CLASSE_ABRAÃO]").__str__()
    alunos.cells[2].text = alunos_matriculados.get("[ALUNOS_MATRICULADOS_CLASSE_SARAH]").__str__()
    alunos.cells[3].text = alunos_matriculados.get("[ALUNOS_MATRICULADOS_CLASSE_AMÓS_E_MIRIÃ]").__str__()
    alunos.cells[4].text = alunos_matriculados.get("[ALUNOS_MATRICULADOS_CLASSE_SAMUEL]").__str__()
    alunos.cells[5].text = alunos_matriculados.get("[ALUNOS_MATRICULADOS_TOTAL_GERAL]").__str__()
    
    ausentes = table.rows[2]
    ausentes.cells[0].text = "Alunos Ausentes"
    ausentes.cells[1].text = alunos_ausentes.get("[AUSENTES_CLASSE_ABRAÃO]").__str__()
    ausentes.cells[2].text = alunos_ausentes.get("[AUSENTES_CLASSE_SARAH]").__str__()
    ausentes.cells[3].text = alunos_ausentes.get("[AUSENTES_CLASSE_AMÓS_E_MIRIÃ]").__str__()
    ausentes.cells[4].text = alunos_ausentes.get("[AUSENTES_CLASSE_SAMUEL]").__str__()
    ausentes.cells[5].text = alunos_ausentes.get("[AUSENTES_TOTAL_GERAL]").__str__()
    
    presentes = table.rows[3]
    presentes.cells[0].text = "Alunos Presentes"
    presentes.cells[1].text = alunos_presentes.get("[PRESENTES_CLASSE_ABRAÃO]").__str__()
    presentes.cells[2].text = alunos_presentes.get("[PRESENTES_CLASSE_SARAH]").__str__()
    presentes.cells[3].text = alunos_presentes.get("[PRESENTES_CLASSE_AMÓS_E_MIRIÃ]").__str__()
    presentes.cells[4].text = alunos_presentes.get("[PRESENTES_CLASSE_SAMUEL]").__str__()
    presentes.cells[5].text = alunos_presentes.get("[PRESENTES_TOTAL_GERAL]").__str__()
    
    r_visitantes = table.rows[4]
    r_visitantes.cells[0].text = "Visitantes"
    r_visitantes.cells[1].text = visitantes.get("[VISITANTES_CLASSE_ABRAÃO]").__str__()
    r_visitantes.cells[2].text = visitantes.get("[VISITANTES_CLASSE_SARAH]").__str__()
    r_visitantes.cells[3].text = visitantes.get("[VISITANTES_CLASSE_AMÓS_E_MIRIÃ]").__str__()
    r_visitantes.cells[4].text = visitantes.get("[VISITANTES_CLASSE_SAMUEL]").__str__()
    r_visitantes.cells[5].text = visitantes.get("[VISITANTES_TOTAL_GERAL]").__str__()
    
    biblias = table.rows[5]
    biblias.cells[0].text = "Bíblias"
    biblias.cells[1].text = quantidade_biblias.get("[BIBLIAS_CLASSE_ABRAÃO]").__str__()
    biblias.cells[2].text = quantidade_biblias.get("[BIBLIAS_CLASSE_SARAH]").__str__()
    biblias.cells[3].text = quantidade_biblias.get("[BIBLIAS_CLASSE_AMÓS_E_MIRIÃ]").__str__()
    biblias.cells[4].text = quantidade_biblias.get("[BIBLIAS_CLASSE_SAMUEL]").__str__()
    biblias.cells[5].text = quantidade_biblias.get("[BIBLIAS_TOTAL_GERAL]").__str__()
    
    revistas = table.rows[6]
    revistas.cells[0].text = "Revistas"
    revistas.cells[1].text = quantidade_revistas.get("[REVISTAS_CLASSE_ABRAÃO]").__str__()
    revistas.cells[2].text = quantidade_revistas.get("[REVISTAS_CLASSE_SARAH]").__str__()
    revistas.cells[3].text = quantidade_revistas.get("[REVISTAS_CLASSE_AMÓS_E_MIRIÃ]").__str__()
    revistas.cells[4].text = quantidade_revistas.get("[REVISTAS_CLASSE_SAMUEL]").__str__()
    revistas.cells[5].text = quantidade_revistas.get("[REVISTAS_TOTAL").__str__()
    
    
 
# Função para estilizar o texto da célula
def set_cell_text_style(cell):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True    
    
def transform_key(key):
    # Exemplo: "[BIBLIAS_CLASSE_ABRAÃO]" -> "Abraão"
    parts = key.split('_')
    nome_classe = ' '.join(parts[2:]).replace('[', '').replace(']', '').title()
    return nome_classe  
    
def ordenar_valores(table):
    sorted_biblias = sorted(quantidade_biblias.items(), key=lambda item: item[1], reverse=True)
    sorted_revistas = sorted(quantidade_revistas.items(), key=lambda item: item[1], reverse=True)
    sorted_presentes = sorted(alunos_presentes.items(), key=lambda item: item[1], reverse=True)
    sorted_visitantes = sorted(visitantes.items(), key=lambda item: item[1], reverse=True)
    
    top_biblias = [transform_key(item[0]) for item in sorted_biblias[:3]]
    top_revistas = [transform_key(item[0]) for item in sorted_revistas[:3]]
    top_presentes = [transform_key(item[0]) for item in sorted_presentes[:3]]
    top_visitantes = [transform_key(item[0]) for item in sorted_visitantes[:3]]
    
    # Preencher a primeira linha (cabeçalho)
    header_cells = table.rows[0].cells
    header_cells[0].text = ""
    header_cells[1].text = "I"
    header_cells[2].text = "II"
    header_cells[3].text = "III"

    # Estilizar o cabeçalho
    for cell in header_cells:
        set_cell_background(cell, '1F3863')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True

    # Preencher a primeira coluna (exceto o cabeçalho)
    first_column = ["", "Bíblias", "Revistas", "Presentes", "Visitantes"]
    for i, text in enumerate(first_column[1:], start=1):
        cell = table.cell(i, 0)
        cell.text = text
        set_cell_background(cell, '1F3863')
        set_cell_text_style(cell)

    # Preencher a tabela com os dados das três primeiras chaves do dicionário ordenado
    for col, key in enumerate(top_biblias, start=1):
        table.rows[1].cells[col].text = key

    for col, key in enumerate(top_revistas, start=1):
        table.rows[2].cells[col].text = key

    for col, key in enumerate(top_presentes, start=1):
        table.rows[3].cells[col].text = key

    for col, key in enumerate(top_visitantes, start=1):
        table.rows[4].cells[col].text = key